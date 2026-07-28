"""
Generates the two output documents used in this division's workflow:
  1. Sanction Memo  (matches "Sanction_Copy" sample format exactly)
  2. NS / Checklist note  (matches "NS_CL_Copy" sample format exactly)

Only the claimant-specific fields change between cases; the surrounding
text is fixed, matching the division's existing wording.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words_inr import rupees_in_words
import datetime


def _set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tc_pr.append(borders)


def _add_para(doc, text, size=11, bold=False, align=None, space_after=6, font='Calibri'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    return p


def build_sanction_memo(case, out_path):
    """
    case: dict with keys -
      memo_no, memo_date, amount, amount_words(optional), name, designation,
      office, pin, emp_id, hospital, hoa (head of account, optional),
      ssp_division (default 'New Delhi Central Division')
    """
    amount_words = case.get('amount_words') or rupees_in_words(case['amount'])
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    _add_para(doc, 'Department of Posts', size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _add_para(doc, f"O/o The Senior Superintendent of Post Offices,", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _add_para(doc, f"{case.get('ssp_division','New Delhi Central Division')}", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _add_para(doc, "Meghdoot Bhawan, New Delhi -110001", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    memo_line = doc.add_paragraph()
    memo_line.paragraph_format.space_after = Pt(14)
    r1 = memo_line.add_run(f"Memo No: {case['memo_no']}")
    r1.font.size = Pt(11)
    r2 = memo_line.add_run(f"\t\t\t\tDated at New Delhi the {case['memo_date']}")
    r2.font.size = Pt(11)

    body_text = (
        f"Sanction of The Sr. Superintendent of Post Offices, {case.get('ssp_division','New Delhi Central Division')}, "
        f"New Delhi 110001 is hereby conveyed for the payment Rs. {case['amount']}/- "
        f"(Rs. {amount_words}) to Sh./Smt. {case['name']}, {case['designation']}, "
        f"{case['office']}, Pin – {case['pin']}, Emp. ID –{case['emp_id']} being the expenses incurred by him/her "
        f"for the treatment/test of {case.get('patient_relation','himself/herself')} from {case['hospital']}, "
        f"under CGHS approved Centre."
    )
    _add_para(doc, body_text, size=11, space_after=14)

    hoa_text = (
        f"The expenditure will be met from the sanction grant under head "
        f"{case.get('hoa', 'GL 3201027306-MR HOA (3201-02-101-01-01-06-MR)')}."
    )
    _add_para(doc, hoa_text, size=11, space_after=28)

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.space_after = 0
    r = sig.add_run("Sr. Supdt. of Post Offices")
    r.font.size = Pt(11)
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.paragraph_format.space_after = 0
    r = sig2.add_run(f"{case.get('ssp_division','New Delhi Central Division')}")
    r.font.size = Pt(11)
    sig3 = doc.add_paragraph()
    sig3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig3.add_run("New Delhi 110001.")
    r.font.size = Pt(11)

    _add_para(doc, "Copy for information and necessary action to:", size=11, space_after=4)
    copy_list = [
        f"1- Sh./Smt. {case['name']}, {case['designation']}, {case['office']}, Pin - {case['pin']}",
        "2- Director, NDHO for information and making payment to the claimant.",
        "3- G M Finance Postal Accounts, Delhi 110054.",
        "4. O/c",
    ]
    for line in copy_list:
        _add_para(doc, line, size=11, space_after=2)

    doc.save(out_path)


def build_ns_checklist(case, items, out_path):
    """
    case: same dict as above, plus:
      cghs_card_no, cghs_validity, submission_date, hco_type ('Govt/CGHS empanelled/Private'),
      nabh_status ('NABH'/'Non-NABH'), treatment_type ('OPD'/'Indoor'/'TEST/INVESTIGATION'/'Emergency')
    items: list of dicts with 'sl_no','particulars','code','claimed','admissible'
    """
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    amount = case['amount']
    amount_words = case.get('amount_words') or rupees_in_words(amount)

    _add_para(doc, case['memo_no'], size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

    intro = (
        f"This is regarding sanction of medical bill of Rs. {amount}/- (Rs. {amount_words}) "
        f"incurred by Sh./Smt. {case['name']}, {case['designation']}, {case['office']}, "
        f"Pin – {case['pin']}, Emp Id- {case['emp_id']} in c/w treatment of "
        f"{case.get('patient_relation','himself/herself')} from, {case['hospital']}, "
        f"which is a CGHS approved lab/hospital."
    )
    _add_para(doc, intro, size=11, space_after=10)

    _add_para(doc, "Admissible amount is as under:-", size=11, bold=True, space_after=6)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    headers = ['Sl. No.', 'Particulars / Code', 'Amount claimed', 'Amount reimbursable']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
        _set_cell_border(hdr[i])

    for it in items:
        row = table.add_row().cells
        row[0].text = str(it.get('sl_no', ''))
        particulars = it.get('particulars', '')
        if it.get('code'):
            particulars += f" ({it['code']})"
        row[1].text = particulars
        row[2].text = f"{it.get('claimed','-')}/-"
        row[3].text = f"{it.get('admissible','-')}/-"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            _set_cell_border(c)

    doc.add_paragraph()
    total_line = (
        f"Total admissible amount comes to Rs. {case.get('total_admissible', amount)}/- "
        f"(Rs. {case.get('total_admissible_words') or rupees_in_words(case.get('total_admissible', amount))} Only). "
        f"In view of above information medical claim case of Rs. {amount}/- is "
        f"{case.get('order_remark', 'in order')}."
    )
    _add_para(doc, total_line, size=11, space_after=8)
    _add_para(doc, f"If agree, we may sanction of Rs. {case.get('total_admissible', amount)}/-", size=11, space_after=8)
    _add_para(doc, "Draft of sanction memo is placed below for approval.", size=11, space_after=16)

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run("Postal Assistant\t\tDySP/ IP(D)/IP (PG)")
    r.font.size = Pt(10.5)

    doc.add_paragraph()
    _add_para(doc, "The check list for the settlement of reimbursement is as under:-", size=11, bold=True, space_after=8)

    checklist_rows = [
        ("1. (i)", "Name & Designation and office of the official", f"Sh./Smt. {case['name']}, {case['designation']}, {case['office']}, Pin – {case['pin']}, Emp Id- {case['emp_id']}"),
        ("(ii)", "Pay and Grade Pay of the official", case.get('pay_grade_pay', '-')),
        ("(iii)", "CGHS card no. of the official, if the official is a CGHS beneficiary and its validity", f"{case.get('cghs_card_no','-')}, {case.get('cghs_validity','-')}"),
        ("(iv)", "Copy of valid CGHS card is placed at", case.get('cghs_card_placed_at', '-')),
        ("2. i)", "Name of the Patient & Relationship with the official", case.get('patient_relation_display', 'SELF')),
        ("ii)", "CGHS Card No. of the patient and its validity", f"{case.get('cghs_card_no','-')}, {case.get('cghs_validity','-')}"),
        ("iii)", "Dependency Certificate (Form-3), if the patient is dependent on the official, is placed at", case.get('dependency_cert', '-')),
        ("3.", "Essentiality Certificate, if applicable is placed at", case.get('essentiality_cert', '-')),
        ("4.", "Medical 97/2004 Form is placed at Sl. No.", case.get('form_97_2004', '-')),
        ("5.", "Whether the treatment is taken in OPD/Indoor/Emergency/with prior permission etc.", case.get('treatment_type', 'TEST/INVESTIGATION')),
        ("7. i)", "Name of the Hospital/Diagnostic Centre from where the treatment is taken", case['hospital']),
        ("ii)", "Whether the hospital/Diagnostic Centre is Govt/CGHS empanelled /Private", case.get('hco_type', 'Yes')),
        ("iii)", "Whether the hospital/Centre is a NABH/Non NABH", case.get('nabh_status', 'NABH')),
        ("8. i)", "Name of the Disease", case.get('disease', '-')),
        ("ii)", "Period of treatment", case.get('period_treatment', '-')),
        ("iii)", "Date of admission", case.get('date_admission', '-')),
        ("iv)", "Date of discharge", case.get('date_discharge', '-')),
        ("v)", "Date of submission of the claim by the claimant", case.get('submission_date', '-')),
        ("13.", "Bills of the hospital/diagnostic centre placed at Sl. No.", case.get('bills_placed_at', '-')),
        ("14.", "Total amount claimed by the official", f"{amount}/"),
        ("15.", "Total admissible amount as per CGHS approved rates", f"{case.get('total_admissible', amount)}/"),
        ("16.", "Amount of advance taken, if any", case.get('advance_taken', '-')),
        ("17", "Any other information related to the case", case.get('other_info', '-')),
        ("18.", "Whether the case is recommended for approval", case.get('recommended', 'Yes')),
    ]

    ctable = doc.add_table(rows=0, cols=3)
    ctable.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.5), Cm(9.5), Cm(6.0)]
    for num, label, val in checklist_rows:
        row = ctable.add_row().cells
        row[0].text = num
        row[1].text = label
        row[2].text = str(val)
        for i, c in enumerate(row):
            c.width = widths[i]
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9.5)
            _set_cell_border(c)

    doc.add_paragraph()
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig2.add_run("Postal Assistant\t\tDySP/ IP(D)/IP (PG)")
    r.font.size = Pt(10.5)

    doc.save(out_path)
